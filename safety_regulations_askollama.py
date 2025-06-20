#safety_regulations_askollama.py
import os, json, traceback, base64, docx, requests
from datetime import datetime
from io import BytesIO
import base64
import gradio as gr
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


OLLAMA_SERVER = "http://127.0.0.1:11434"
SAVE_DIR = "uploaded_data" #store user uploaded data
REPORT_DIR = "user_reports" #store user reports
SYSTEM_VERSION = "v1.0.0"
REQUEST_TIMEOUT = 200  # 3 minutes timeout for API requests
MODEL_USED = "llava:13b"
#ALL MODELS: CHANGE THE NAME 
# llava:7b, bakllava:7b, gemma3:12b, llava:13b, qwen2.5vl:7b


# Ensure the directory exists
os.makedirs(SAVE_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

def parse_markdown_table(md_text):
    """
    Parse a Markdown table string into a list of rows (each row is a list of cell strings).
    Assumes table uses pipes and has a header separator line.
    Returns (rows, original_table_text).
    """
    lines = md_text.splitlines()
    table_lines = []
    rows = []
    in_table = False
    for line in lines:
        if '|' in line:
            in_table = True
            table_lines.append(line)
            if set(line.strip()) <= set("|- :"):
                continue
            parts = [cell.strip() for cell in line.split('|')]
            if parts and parts[0] == '': parts = parts[1:]
            if parts and parts[-1] == '': parts = parts[:-1]
            rows.append(parts)
        elif in_table:
            # table ended
            break
    original_table = "\n".join(table_lines)
    return rows, original_table


def prepare_image(image):
    """Prepare image for API submission"""
   # Convert to RGB if needed
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    # Convert to base64 with optimized format
    buf = BytesIO()
    image.save(buf, format="JPEG", quality=85, optimize=True)
    return buf.getvalue()

def format_error(message):
    return f"<span style='color: red;'>❗ {message}</span>"

def ask_ollama(image, prompt, input_name, input_staff_id, input_sector, run_analysis=False):
    """
    Sends image+prompt to Ollama, saves input, parses the table output,
    saves a JSON with headers and creates a Word document
    with a proper table and returns both Markdown and .docx path.
    """
    
    # Validate all inputs
    if image is None:
        return format_error("Oops! You forgot to add a photo 📸"), None, None
    if not all([input_name, input_staff_id, input_sector]):
        return format_error("Wait! We still need your deets first 👤"), None, None
    if not run_analysis:
        # Show loading message before running analysis
        return (
            "<div style='text-align: center; padding: 20px; font-size: 18px;'>"
            "<div style='margin-bottom: 10px; font-size: 22px;'>🔄 Working on it....</div>"
            "<div style='color: #666; font-size: 16px;'>Grab a coffee, this might take a minute ☕</div>"
            "</div>",
            None,
            ""
        )
    # Proceed with analysis
    try:
        image_data = prepare_image(image)
        if image_data is None:
            return format_error("Error processing image."), None, None

        # Convert to base64
        img_b64 = base64.b64encode(image_data).decode()

        # Build model instruction
        instruction = (
            "**Role:** You are an expert in occupational safety and health. "
            "**Objective:** Your task is to analyze the uploaded image for potential workplace hazards based on the DOSH Malaysia, Akta Keselamatan Dan Kesihatan Pekerjaan (Pindaan) 2022 (Akta A1648).\n\n"

            "### Response Format Requirements:\n"
            "1. Output MUST be a single Markdown table with exactly 6 columns.\n"
            "2. NO text before or after the table.\n"
            "3. NO explanations or additional content.\n\n"

            "### Table Structure:\n"
            "| No. | Work Activity | Possible Hazard | Existing Risk Control (if any) | Severity | Recommended Control Measures |\n"
            "|-----|--------------|-----------------|--------------------------------|----------|------------------------------|\n\n"

            "### Column Guidelines:\n"
            "1. **No.**: Sequential numbers (1, 2, 3...)\n"
            "2. **Work Activity**: Brief description of the activity shown in the image\n"
            "3. **Possible Hazard**: Specific hazard identified in the activity\n"
            "4. **Existing Risk Control**: Write 'No' if none exists, otherwise brief description\n"
            "5. **Severity**: ONLY use 'High', 'Moderate', or 'Low'\n"
            "6. **Recommended Control Measures**: Specific safety measures to implement\n\n"

            "### Content Rules:\n"
            "1. Sort rows by Severity: High → Moderate → Low\n"
            "2. Each hazard must be specific to the image\n"
            "3. Use concise, clear language\n"
            "4. Avoid repeating the same Work Activity unless for different hazards\n"
            "5. Each cell should be brief and to the point\n\n"

            "### Example Row:\n"
            "| 1 | Welding in confined space | Oxygen deficiency | No | High | Ensure proper ventilation and use gas detectors |\n\n"

            "### Important Notes:\n"
            "1. Focus ONLY on hazards visible in the image\n"
            "2. Base recommendations on DOSH Malaysia regulations\n"
            "3. Be specific and practical in control measures\n"
            "4. Maintain consistent formatting throughout\n"
        )
        full_prompt = instruction + "\n### User Prompt (Context):\n" + prompt
        
        # Payload for inference
        payload = {
            "model": MODEL_USED,
            "prompt": full_prompt,
            "stream": False,
            "images": [img_b64]
        }

        # Start timing
        start_time = datetime.now()
        
        # Make API request with timeout
        resp = requests.post(f"{OLLAMA_SERVER}/api/generate", json=payload, timeout=REQUEST_TIMEOUT)
        resp.raise_for_status()
        text = resp.json().get('response', '')
        
        # Calculate generation time
        end_time = datetime.now()
        generation_time = (end_time - start_time).total_seconds()

        # Parse table and extract original table text
        rows_list, _ = parse_markdown_table(text)
        headers = rows_list[0] if rows_list else []

        # Store data for later use
        analysis_data = {
            "text": text,
            "rows_list": rows_list,
            "headers": headers,
            "generation_time": generation_time,
            "image": image,
            "prompt": prompt,
            "input_name": input_name,
            "input_staff_id": input_staff_id,
            "input_sector": input_sector,
            "model_used": payload["model"]
        }

        return text, analysis_data, None

    except requests.Timeout:
        return format_error("Request timed out. Please try again."), None, None
    except requests.RequestException as e:
        return format_error(f"API request failed: {str(e)}"), None, None
    except Exception as e:
        error_details = traceback.format_exc()
        print(error_details)  # For logging
        return format_error(f"Error: {str(e)}"), None, None
    
def save_analysis(analysis_data):
    """
    Save the analysis to JSON and Word files after user confirmation
    """
    if not analysis_data:
        return "No analysis data to save.", None

    try:
        # Timestamped filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{timestamp}"

        # Save image
        img_path = os.path.join(SAVE_DIR, f"{base_filename}_image.png")
        analysis_data["image"].save(img_path)

        # Build JSON output
        json_output = {
            "timestamp": timestamp,
            "model_used": analysis_data["model_used"],
            "generation_time_seconds": round(analysis_data["generation_time"], 2),
            "name": analysis_data["input_name"],
            "staff_id": analysis_data["input_staff_id"],
            "sector": analysis_data["input_sector"],
            "image_path": img_path,
            "system_version": SYSTEM_VERSION,
            "prompt": analysis_data["prompt"],
            "table_headers": analysis_data["headers"],
            "rows": analysis_data["rows_list"][1:]
        }

        # Save JSON
        json_path = os.path.join(SAVE_DIR, f"{base_filename}.json")
        with open(json_path, 'w') as jf:
            json.dump(json_output, jf, indent=4)

        # Build Word doc
        doc = Document()
        
        # HEADING
        heading = doc.add_heading('OCCUPATIONAL SAFETY AND HEALTH RESULT', level=1)
        heading.alignment = 1  # Center alignment (0=left, 1=center, 2=right)
        for run in heading.runs:
            run.bold = True
        doc.add_paragraph()
        
        # USER INFO TABLE
        user_info_table = doc.add_table(rows=2, cols=4)
        user_info_table.style = 'Table Grid'
        user_info_table.autofit = False

        # Set column widths (in inches)
        for row in user_info_table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(1.5)

        # Add user information to the table
        display_date = datetime.strptime(timestamp, "%Y%m%d_%H%M%S").strftime("%d-%m-%Y") # Format date format (DD-MM-YYYY)
        # First row (labels)
        labels = ["NAME", "STAFF ID", "SECTOR", "DATE"]
        for i, label in enumerate(labels):
            cell = user_info_table.cell(0, i)
            cell.text = label
            # Make the text bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Second row (values)
        values = [analysis_data["input_name"], analysis_data["input_staff_id"], analysis_data["input_sector"], display_date]
        for i, value in enumerate(values):
            user_info_table.cell(1, i).text = value
        # Style the text
        for row in user_info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(11)
        doc.add_paragraph()
        
        # IMAGE AND PROMPT
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_path, width=docx.shared.Inches(3.0))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = doc.add_paragraph()
        p.add_run('PROMPT: ').bold = True
        doc.add_paragraph()
        p.add_run(f'"{analysis_data["prompt"]}"')
        
        doc.add_page_break()
        # TABLE
        tbl = doc.add_table(rows=1, cols=len(analysis_data["headers"]))
        tbl.style = 'Table Grid'
        tbl.autofit = False

        column_widths = {
            "No.": 0.4,
            "Work Activity": 1.0,
            "Possible Hazard": 0.8,
            "Existing Risk Control (if any)": 1.0,
            "Severity": 0.8,
            "Recommended Control Measures": 2.4

        }
        # Apply column widths
        for idx, header in enumerate(analysis_data["headers"]):
            cell = tbl.rows[0].cells[idx]
            cell.text = header
            cell.width = Inches(column_widths.get(header, 1.0))
            # Format header
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1  # Center alignment
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = docx.shared.Pt(11)
        # Add data rows
        for row in analysis_data["rows_list"][1:]:
            rc = tbl.add_row().cells
            for i, cell_text in enumerate(row):
                rc[i].text = cell_text
                rc[i].width = Inches(list(column_widths.values())[i])
                # Format cell content
                for paragraph in rc[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(10)

        docx_path = os.path.join(SAVE_DIR, f"{base_filename}_analysis.docx")
        doc.save(docx_path)

        return (
            "<span style='font-weight: bold;'>✅ Files saved successfully!</span> <br>"
            "You can now download your report above.<br>",
            docx_path
        )

    except Exception as e:
        error_details = traceback.format_exc()
        print(error_details)  # For logging
        return f"Error saving files: {str(e)}", None

def save_report(analysis_data, email, description, rating):
    if not all([description, rating]):
        return format_error("Please fill all required fields!"), "", ""
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"{timestamp}"

    name = analysis_data["input_name"] if analysis_data and "input_name" in analysis_data else "NA"
    staff_id = analysis_data["input_staff_id"] if analysis_data and "input_staff_id" in analysis_data else "NA"
    sector = analysis_data["input_sector"] if analysis_data and "input_sector" in analysis_data else "NA"
    model_used = analysis_data["model_used"] if analysis_data and "model_used" in analysis_data else "NA"
    email = email if email else "NA"

    report_data = {
        "timestamp": timestamp,
        "name": name,
        "staff_id": staff_id,
        "sector": sector,
        "model_used": model_used,
        "email": email,
        "description": description.strip(),
        "rating": int(rating),
        "system_version": SYSTEM_VERSION,
    }
    try:
        report_path = os.path.join(REPORT_DIR, f"{base_filename}_report.json")
        with open(report_path, "w") as f:
            json.dump(report_data, f, indent=2)
        return "<span style='color: green;'>Thank You! We appreciate your response!</span>", "", ""
    except Exception as e:
        return format_error(f"⚠️ Oops! Failed to save report: {str(e)}"), None, None

# INTERFACE DESIGN
with gr.Blocks(title="Occupational Safety and Health Chatbot") as iface:
    gr.Markdown("""
        <div style='display: flex; align-items: center; justify-content: center; gap: 15px; margin-bottom: 20px;'>
            <div style='font-size: 24px; font-weight: bold;'>OCCUPATIONAL SAFETY AND HEALTH CHATBOT</div>
        </div>
        """)
    with gr.Row():
        input_name = gr.Textbox(placeholder="Enter your name", label="NAME")
        input_staff_id = gr.Textbox(placeholder="Enter your staff ID", label="STAFF ID")
        input_sector = gr.Dropdown(
            choices=["CEO & BOD Office", "Governance & Operations Sector", "Technology Advocacy Sector", "Technology Development Sector", "Technology Venture Sector", "MSSB", "MTSBB"],
            label="SECTOR",
            value=None,
            allow_custom_value=True
        )

    gr.Markdown("*Upload an image of a workplace and **_(OPTIONALLY)_** describe the activity of the image.*")
    with gr.Row():
        with gr.Column(scale=1):  # This will take 1/4 of the row
            input_image = gr.Image(
                type="pil", 
                label="Upload Image",
                height=300,  # Fixed height in pixels
                width=300,   # Fixed width in pixels
                image_mode="RGB"
            )
            input_prompt = gr.Textbox(lines=2, placeholder="Enter your prompt", label="IMAGE DESCRIPTION", info="💡*Example prompt: 'Welding work in confined space'*")
            submit_btn = gr.Button("Start")

        with gr.Column(scale=3):  # This will take 3/4 of the row
            output_markdown = gr.Markdown(label="Analysis (Markdown)")
            export_btn = gr.Button("Export Report Analysis (.docx)", visible=False)
            download_btn = gr.DownloadButton(label="Download Safety Report", visible=False, variant="primary")
            status_msg = gr.Markdown()
            output_file = gr.File(label="Safety Report", visible=False)

    # Store analysis data
    analysis_data = gr.State(None)

    def show_export_button(analysis_data):
        if analysis_data:
            return gr.update(visible=True)
        return gr.update(visible=False)


    # Connect the buttons
    submit_btn.click(
        fn=lambda image, prompt, name, staff_id, sector: ask_ollama(image, prompt, name, staff_id, sector, False),
        inputs=[input_image, input_prompt, input_name, input_staff_id, input_sector],
        outputs=[output_markdown, analysis_data, status_msg]
    ).then(
        fn=lambda image, prompt, name, staff_id, sector: ask_ollama(image, prompt, name, staff_id, sector, True),
        inputs=[input_image, input_prompt, input_name, input_staff_id, input_sector],
        outputs=[output_markdown, analysis_data, status_msg]
    ).then(
        fn=show_export_button,
        inputs=[analysis_data],
        outputs=[export_btn]
    )

    export_btn.click(
        fn=lambda: ("<span style='color: #666;'>Generating document...</span>", None),
        inputs=[],
        outputs=[status_msg, output_file]
    ).then(
        fn=save_analysis,
        inputs=[analysis_data],
        outputs=[status_msg, output_file]
    ).then(
        fn=lambda file: gr.update(visible=True),
        inputs=[output_file],
        outputs=[download_btn]
    )

    download_btn.click(
        fn=lambda file: file,
        inputs=[output_file],
        outputs=[download_btn]
    )

    # Place this after the main analysis UI
    with gr.Accordion("SHARE YOUR THOUGHTS", open=False):
        with gr.Row():
            report_email = gr.Textbox(label="Email Address (optional)", placeholder="user@example.com")
            report_rating = gr.Slider(
                minimum=1, 
                maximum=5, 
                step=1, 
                value=5, 
                label="Rate your experience",
                info="1=Meh, 5=Awesome"
            )
        report_issue = gr.Textbox(label="Experience or Issue", lines=2, placeholder="Share with us"
        )
        report_status = gr.Markdown()
        report_btn = gr.Button("Submit")

    report_btn.click(
        fn=save_report,
        inputs=[analysis_data, report_email, report_issue, report_rating],
        outputs=[report_status, report_email, report_issue]
    )

if __name__ == '__main__':
    iface.launch(
        server_name='0.0.0.0',
        server_port=8002,
        share=True
    )
